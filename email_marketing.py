"""
MODULE EMAIL MARKETING ABED ACADEMY - Version 3
=================================================
- Tous les parcours ABED détaillés
- Filtre junior max 3 ans
- Envoi via Resend
"""

import os
import json
import base64
import logging
from datetime import date, datetime
from pathlib import Path

import anthropic
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import resend

log = logging.getLogger("veille")

PARCOURS_ABED = [
    {
        "nom": "Coaching à l'Entrepreneuriat Agricole Innovant et Inclusif",
        "categorie": "Accompagnement à l'Entrepreneuriat Agricole",
        "url": "https://academy.abedong.org/course/view.php?id=28",
        "competences": ["entrepreneuriat agricole","agribusiness","coaching agricole",
                        "agriculture","développement rural","chaîne de valeur",
                        "projet agricole","innovation agricole","incubation","idéation",
                        "accompagnement","jeunes entrepreneurs"],
        "metiers_cibles": ["conseiller agricole","coach agricole","chargé de projet agricole",
                           "coordinateur agri","responsable filière","animateur rural",
                           "technicien agricole"],
        "description": "Forme à l'accompagnement en entrepreneuriat agricole, de la maturation des idées à la mise en œuvre.",
    },
    {
        "nom": "Conseiller Junior en Education Financière",
        "categorie": "Finances & Gestion",
        "url": "https://academy.abedong.org/course/index.php?categoryid=8",
        "competences": ["éducation financière","finance personnelle","conseil financier",
                        "épargne","budget","microfinance","inclusion financière"],
        "metiers_cibles": ["conseiller financier","agent de crédit","chargé d'éducation financière",
                           "animateur financier","conseiller microfinance"],
        "description": "Forme des conseillers capables d'accompagner ménages et petites entreprises dans leur gestion.",
    },
    {
        "nom": "Assistant Administratif et Financier",
        "categorie": "Finances & Gestion",
        "url": "https://academy.abedong.org/course/index.php?categoryid=8",
        "competences": ["administration","finance","gestion administrative","comptabilité de base",
                        "classement","archivage","gestion de caisse","reporting",
                        "assistance administrative","secrétariat"],
        "metiers_cibles": ["assistant administratif","assistant financier","secrétaire comptable",
                           "aide administratif","agent administratif"],
        "description": "Prépare aux tâches administratives et financières au sein d'entreprises, ONG et projets.",
    },
    {
        "nom": "Comptable d'Entreprise",
        "categorie": "Finances & Gestion",
        "url": "https://academy.abedong.org/course/index.php?categoryid=8",
        "competences": ["comptabilité","comptabilité générale","bilan","SYSCOHADA",
                        "journal","grand livre","déclaration fiscale","TVA","SAGE"],
        "metiers_cibles": ["comptable","aide-comptable","comptable junior","agent comptable",
                           "technicien comptable"],
        "description": "Forme aux fondamentaux de la comptabilité d'entreprise selon les normes SYSCOHADA.",
    },
    {
        "nom": "Comptable de Projet",
        "categorie": "Finances & Gestion",
        "url": "https://academy.abedong.org/course/index.php?categoryid=8",
        "competences": ["comptabilité de projet","gestion budgétaire","rapport financier",
                        "suivi budgétaire","bailleurs de fonds","ONG",
                        "projet de développement","plan de travail budgétisé"],
        "metiers_cibles": ["comptable de projet","gestionnaire financier de projet",
                           "assistant financier ONG","chargé de finances projet"],
        "description": "Spécialisé dans la gestion financière des projets de développement financés par bailleurs.",
    },
    {
        "nom": "Assistant en Suivi-Evaluation (MEAL)",
        "categorie": "Suivi-Evaluation",
        "url": "https://academy.abedong.org/course/view.php?id=16",
        "competences": ["suivi-évaluation","MEAL","monitoring","évaluation","cadre logique",
                        "indicateurs","collecte de données","KoBoToolbox","ODK",
                        "analyse de données","tableau de bord","redevabilité"],
        "metiers_cibles": ["assistant suivi-évaluation","chargé de suivi-évaluation",
                           "responsable MEAL","coordinateur M&E","chargé de programme",
                           "agent de collecte de données"],
        "description": "Forme au métier d'assistant suivi-évaluation MEAL pour projets de développement et ONG.",
    },
    {
        "nom": "Animateur en Nutrition Communautaire",
        "categorie": "Nutrition Communautaire & Transformation Agroalimentaire",
        "url": "https://academy.abedong.org/course/index.php?categoryid=7",
        "competences": ["nutrition communautaire","santé communautaire","ANJE",
                        "malnutrition","sensibilisation nutritionnelle","allaitement",
                        "dépistage nutritionnel","MUAC","animation communautaire"],
        "metiers_cibles": ["animateur nutrition","agent de santé communautaire",
                           "relais communautaire","agent WASH","chargé de nutrition"],
        "description": "Prépare à l'animation et à la sensibilisation nutritionnelle au niveau communautaire.",
    },
    {
        "nom": "Conseiller Technique en Transformation Agroalimentaire",
        "categorie": "Nutrition Communautaire & Transformation Agroalimentaire",
        "url": "https://academy.abedong.org/course/index.php?categoryid=7",
        "competences": ["transformation agroalimentaire","qualité alimentaire","HACCP",
                        "hygiène alimentaire","conservation des aliments",
                        "conditionnement","agro-industrie","valorisation des produits"],
        "metiers_cibles": ["technicien agroalimentaire","agent qualité alimentaire",
                           "conseiller en transformation","technicien qualité"],
        "description": "Forme aux techniques de transformation et qualité dans les filières agroalimentaires.",
    },
    {
        "nom": "Conseil Agricole en Agroécologie",
        "categorie": "Conseils Agricoles",
        "url": "https://academy.abedong.org/course/index.php?categoryid=11",
        "competences": ["agroécologie","agriculture durable","compostage","permaculture",
                        "biodiversité","agroforesterie","gestion intégrée",
                        "conseil agroécologique","transition agro-écologique"],
        "metiers_cibles": ["conseiller agroécologie","agent de vulgarisation agroécologie",
                           "technicien agroécologie","animateur agriculture durable"],
        "description": "Spécialisé dans le conseil et l'accompagnement des producteurs vers l'agroécologie.",
    },
    {
        "nom": "Conseil Agricole Général en Contexte Africain",
        "categorie": "Conseils Agricoles",
        "url": "https://academy.abedong.org/course/index.php?categoryid=11",
        "competences": ["conseil agricole","vulgarisation agricole","appui technique",
                        "diagnostic agronomique","systèmes de production","cultures vivrières",
                        "producteurs","coopératives","itinéraire technique"],
        "metiers_cibles": ["conseiller agricole","agent de vulgarisation","technicien agricole",
                           "responsable formation agricole","agent d'appui technique"],
        "description": "Forme des conseillers agricoles polyvalents adaptés aux réalités des exploitations africaines.",
    },
]


def match_parcours(offre: dict) -> dict:
    texte = (offre.get("titre","") + " " + offre.get("resume","") + " " +
             offre.get("secteur","") + " " + offre.get("raw_text","")).lower()
    best_parcours = PARCOURS_ABED[0]
    best_score    = 0
    for p in PARCOURS_ABED:
        score  = sum(2 for m in p["competences"] if m.lower() in texte)
        score += sum(3 for m in p["metiers_cibles"] if m.lower() in texte)
        if score > best_score:
            best_score    = score
            best_parcours = p
    return best_parcours


def filter_junior_offers(offers: list[dict]) -> list[dict]:
    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY",""))
    items  = json.dumps([
        {"id":i,"titre":o.get("titre",""),"raw_text":o.get("raw_text","")[:300]}
        for i,o in enumerate(offers)
    ], ensure_ascii=False)
    prompt = f"""Analyse ces offres et détermine l'expérience requise.
{items}
Pour chaque offre JSON :
- id
- annees_experience : entier (0 si débutant, sinon nombre d'années)
- mention_experience : texte exact trouvé ou "non précisé"
- accessible_junior : true si expérience <= 3 ans OU non précisée
Réponds UNIQUEMENT tableau JSON valide."""
    try:
        resp = client.messages.create(
            model="claude-sonnet-4-5", max_tokens=2000,
            messages=[{"role":"user","content":prompt}]
        )
        raw = resp.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"): raw = raw[4:]
        results  = json.loads(raw)
        filtered = []
        excluded = 0
        for item in results:
            idx = item.get("id",0)
            if idx < len(offers):
                o = dict(offers[idx])
                o["annees_experience"]  = item.get("annees_experience",0)
                o["mention_experience"] = item.get("mention_experience","non précisé")
                o["accessible_junior"]  = item.get("accessible_junior",True)
                if o["accessible_junior"]:
                    filtered.append(o)
                else:
                    excluded += 1
        log.info(f"  Filtre junior : {len(filtered)} gardées, {excluded} exclues (>3 ans)")
        return filtered
    except Exception as e:
        log.error(f"Erreur filtre junior: {e}")
        for o in offers:
            o["annees_experience"]  = 0
            o["mention_experience"] = "non précisé"
            o["accessible_junior"]  = True
        return offers


def select_top10_for_abed(offers: list[dict]) -> list[dict]:
    log.info("Filtrage offres junior (max 3 ans exp)...")
    junior_offers = filter_junior_offers(offers)
    if not junior_offers:
        log.warning("Aucune offre junior — on garde toutes")
        junior_offers = offers

    scored = []
    for o in junior_offers:
        parcours = match_parcours(o)
        texte    = (o.get("titre","") + " " + o.get("raw_text","")).lower()
        score    = sum(2 for m in parcours["competences"] if m.lower() in texte)
        score   += sum(3 for m in parcours["metiers_cibles"] if m.lower() in texte)
        score   += o.get("pertinence_score",3)
        scored.append((score, o, parcours))
    scored.sort(key=lambda x: -x[0])

    top10 = []
    parcours_utilises = {}
    for score, offre, parcours in scored:
        nom_p = parcours["nom"]
        if parcours_utilises.get(nom_p,0) < 2:
            offre["_parcours_matche"] = parcours
            top10.append(offre)
            parcours_utilises[nom_p] = parcours_utilises.get(nom_p,0) + 1
        if len(top10) >= 10: break

    if len(top10) < 10:
        for score, offre, parcours in scored:
            if offre not in top10:
                offre["_parcours_matche"] = parcours
                top10.append(offre)
            if len(top10) >= 10: break

    return top10


def enrich_offers_for_marketing(offers: list[dict]) -> list[dict]:
    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY",""))
    parcours_ctx = json.dumps([
        {"nom":p["nom"],"url":p["url"],"description":p["description"],
         "competences":p["competences"][:5]}
        for p in PARCOURS_ABED
    ], ensure_ascii=False)
    items = json.dumps([
        {"id":i,"titre":o.get("titre",""),"org":o.get("org",""),
         "pays":o.get("pays",""),"secteur":o.get("secteur",""),
         "resume":o.get("resume",""),"experience":o.get("mention_experience","?"),
         "parcours_abed":o.get("_parcours_matche",{}).get("nom","")}
        for i,o in enumerate(offers)
    ], ensure_ascii=False)
    prompt = f"""Tu es directeur marketing ABED Academy (Bénin).
Parcours : {parcours_ctx}
Offres : {items}
Pour CHAQUE offre JSON :
- id
- competences_requises : liste 3-5 compétences clés
- argumentaire_abed : 2 phrases percutantes (résultats concrets)
- niveau_requis : "Débutant", "Junior (1-2 ans)", ou "Confirmé (2-3 ans)"
- conseil_candidature : 1 conseil pratique
Réponds UNIQUEMENT tableau JSON valide."""
    try:
        resp = client.messages.create(
            model="claude-sonnet-4-5", max_tokens=3000,
            messages=[{"role":"user","content":prompt}]
        )
        raw = resp.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"): raw = raw[4:]
        for item in json.loads(raw):
            idx = item.get("id",0)
            if idx < len(offers):
                offers[idx]["competences_requises"] = item.get("competences_requises",[])
                offers[idx]["argumentaire_abed"]     = item.get("argumentaire_abed","")
                offers[idx]["niveau_requis"]          = item.get("niveau_requis","Junior")
                offers[idx]["conseil_candidature"]    = item.get("conseil_candidature","")
        log.info("Enrichissement marketing terminé")
    except Exception as e:
        log.error(f"Erreur enrichissement: {e}")
        for o in offers:
            if "argumentaire_abed" not in o:
                p = o.get("_parcours_matche",{})
                o["competences_requises"] = []
                o["argumentaire_abed"]     = f"Le parcours '{p.get('nom','')}' vous prépare directement à ce poste."
                o["niveau_requis"]          = "Junior"
                o["conseil_candidature"]    = "Mettez en avant votre formation ABED Academy dans votre CV."
    return offers


def generate_excel(offers: list[dict], output_path: Path) -> Path:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Top 10 Offres Junior"
    VERT  = "0F6E56"; VCLAIR = "E1F5EE"; BLANC = "FFFFFF"; GRIS = "F8F7F2"; ORANGE = "EF9F27"
    bd = Border(left=Side(style='thin',color="CCCCCC"),right=Side(style='thin',color="CCCCCC"),
                top=Side(style='thin',color="CCCCCC"),bottom=Side(style='thin',color="CCCCCC"))

    ws.merge_cells("A1:I1")
    ws["A1"] = f"ABED ACADEMY — Top 10 Offres Junior | {date.today().strftime('%d %B %Y')}"
    ws["A1"].font      = Font(name="Arial",size=14,bold=True,color=BLANC)
    ws["A1"].fill      = PatternFill("solid",fgColor=VERT)
    ws["A1"].alignment = Alignment(horizontal="center",vertical="center")
    ws.row_dimensions[1].height = 32

    ws.merge_cells("A2:I2")
    ws["A2"] = "Offres accessibles aux juniors (max 3 ans) — liées aux parcours ABED Academy | academy.abedong.org"
    ws["A2"].font      = Font(name="Arial",size=10,italic=True,color=VERT)
    ws["A2"].fill      = PatternFill("solid",fgColor=VCLAIR)
    ws["A2"].alignment = Alignment(horizontal="center",vertical="center")
    ws.row_dimensions[2].height = 20

    headers = ["N°","Titre du Poste","Organisation","Pays","Expérience requise",
               "Parcours ABED Recommandé","Compétences Requises","Argumentaire ABED Academy","Lien pour Postuler"]
    for col,h in enumerate(headers,1):
        cell = ws.cell(row=4,column=col,value=h)
        cell.font      = Font(name="Arial",size=10,bold=True,color=BLANC)
        cell.fill      = PatternFill("solid",fgColor=VERT)
        cell.alignment = Alignment(horizontal="center",vertical="center",wrap_text=True)
        cell.border    = bd
    ws.row_dimensions[4].height = 28

    for i,o in enumerate(offers):
        row      = i+5
        parcours = o.get("_parcours_matche",{})
        comp     = "\n".join(f"• {c}" for c in o.get("competences_requises",[]))
        lien     = o.get("url","") or "Non disponible"
        bg       = BLANC if i%2==0 else GRIS
        values   = [i+1,o.get("titre",""),o.get("org",""),o.get("pays",""),
                    o.get("mention_experience","Non précisé"),parcours.get("nom",""),
                    comp,o.get("argumentaire_abed",""),lien]
        for col,val in enumerate(values,1):
            cell = ws.cell(row=row,column=col,value=val)
            cell.font      = Font(name="Arial",size=9)
            cell.fill      = PatternFill("solid",fgColor=bg)
            cell.border    = bd
            cell.alignment = Alignment(vertical="top",wrap_text=True)
            if col==1: cell.font=Font(name="Arial",size=10,bold=True,color=VERT); cell.alignment=Alignment(horizontal="center",vertical="center")
            if col==2: cell.font=Font(name="Arial",size=9,bold=True)
            if col==5: cell.font=Font(name="Arial",size=9,color=ORANGE,bold=True)
            if col==6: cell.font=Font(name="Arial",size=9,color=VERT,bold=True)
            if col==9 and lien!="Non disponible": cell.font=Font(name="Arial",size=9,color="185FA5",underline="single"); cell.hyperlink=lien
        ws.row_dimensions[row].height = 85

    for col,w in enumerate([4,32,26,12,18,30,25,42,35],1):
        ws.column_dimensions[get_column_letter(col)].width = w

    ws2 = wb.create_sheet("Parcours ABED Academy")
    ws2.merge_cells("A1:D1")
    ws2["A1"] = "Parcours ABED Academy — Liens d'inscription"
    ws2["A1"].font=Font(name="Arial",size=13,bold=True,color=BLANC)
    ws2["A1"].fill=PatternFill("solid",fgColor=VERT)
    ws2["A1"].alignment=Alignment(horizontal="center",vertical="center")
    for col,h in enumerate(["Parcours","Catégorie","Description","Lien"],1):
        cell=ws2.cell(row=2,column=col,value=h)
        cell.font=Font(name="Arial",size=10,bold=True,color=BLANC)
        cell.fill=PatternFill("solid",fgColor=VERT); cell.border=bd
    for i,p in enumerate(PARCOURS_ABED):
        row=i+3; bg=BLANC if i%2==0 else GRIS
        for col,val in enumerate([p["nom"],p["categorie"],p["description"],p["url"]],1):
            cell=ws2.cell(row=row,column=col,value=val)
            cell.font=Font(name="Arial",size=9,color="185FA5" if col==4 else "000000",underline="single" if col==4 else None)
            cell.fill=PatternFill("solid",fgColor=bg); cell.border=bd
            cell.alignment=Alignment(vertical="top",wrap_text=True)
            if col==4: cell.hyperlink=val
        ws2.row_dimensions[row].height=30
    for col,w in enumerate([40,35,50,50],1):
        ws2.column_dimensions[get_column_letter(col)].width=w

    wb.save(output_path)
    log.info(f"Excel généré : {output_path}")
    return output_path


def generate_word(offers: list[dict], output_path: Path) -> Path:
    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY",""))
    offres_ctx = json.dumps([
        {"titre":o.get("titre",""),"org":o.get("org",""),"pays":o.get("pays",""),
         "secteur":o.get("secteur",""),"resume":o.get("resume",""),
         "experience":o.get("mention_experience","non précisé"),
         "parcours_abed":o.get("_parcours_matche",{}).get("nom",""),
         "url_parcours":o.get("_parcours_matche",{}).get("url",""),
         "argumentaire":o.get("argumentaire_abed","")}
        for o in offers[:5]
    ], ensure_ascii=False)

    posts_linkedin = []
    scripts_tiktok = []

    try:
        r1 = client.messages.create(model="claude-sonnet-4-5",max_tokens=3000,
            messages=[{"role":"user","content":
                f"""Community manager ABED Academy (academy.abedong.org), Bénin.
Top 5 offres junior du jour : {offres_ctx}
Génère 5 posts LinkedIn (150-200 mots chacun) :
- Accroche forte (stat ou question choc)
- Cite l'offre réelle comme preuve de marché
- Présente le parcours ABED comme solution pour juniors (0-3 ans exp)
- CTA vers academy.abedong.org
- 5-7 hashtags (#emploibenin #abedacademy #formation etc.)
Format JSON : [{{"titre":"...","contenu":"...","hashtags":"..."}}]
UNIQUEMENT JSON valide."""}])
        raw1=r1.content[0].text.strip()
        if raw1.startswith("```"): raw1=raw1.split("```")[1]; raw1=raw1[4:] if raw1.startswith("json") else raw1
        posts_linkedin=json.loads(raw1)
        log.info(f"  {len(posts_linkedin)} posts LinkedIn générés")
    except Exception as e:
        log.error(f"Erreur LinkedIn: {e}")

    try:
        r2 = client.messages.create(model="claude-sonnet-4-5",max_tokens=2000,
            messages=[{"role":"user","content":
                f"""Créateur TikTok ABED Academy (academy.abedong.org).
Offres : {offres_ctx}
Génère 2 scripts TikTok (30-60 sec) :
- Accroche visuelle choc 3 secondes
- Cite offres béninoises/africaines réelles
- Montre que juniors peuvent décrocher ces postes avec ABED
- CTA academy.abedong.org
Format JSON : [{{"titre":"...","accroche_visuelle":"...","script_parle":"...","texte_ecran":"...","hashtags":"..."}}]
UNIQUEMENT JSON valide."""}])
        raw2=r2.content[0].text.strip()
        if raw2.startswith("```"): raw2=raw2.split("```")[1]; raw2=raw2[4:] if raw2.startswith("json") else raw2
        scripts_tiktok=json.loads(raw2)
        log.info(f"  {len(scripts_tiktok)} scripts TikTok générés")
    except Exception as e:
        log.error(f"Erreur TikTok: {e}")

    doc = Document()
    VERT=RGBColor(0x0F,0x6E,0x56); VERT2=RGBColor(0x1D,0x9E,0x75)
    GRIS=RGBColor(0x5F,0x5E,0x5A); ORANGE=RGBColor(0xEF,0x9F,0x27)

    def h(text,size=16,color=None,bold=True):
        p=doc.add_paragraph(); r=p.add_run(text)
        r.bold=bold; r.font.name="Arial"; r.font.size=Pt(size)
        r.font.color.rgb=color or VERT
        p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(6)
        return p

    def sep():
        p=doc.add_paragraph(); r=p.add_run("─"*80)
        r.font.size=Pt(8); r.font.color.rgb=RGBColor(0xCC,0xCC,0xCC)

    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("ABED ACADEMY"); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=VERT

    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=p2.add_run("Contenu Marketing Quotidien — Offres Junior (0-3 ans)")
    r2.font.size=Pt(14); r2.font.color.rgb=GRIS

    p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r3=p3.add_run(date.today().strftime('%A %d %B %Y').capitalize())
    r3.italic=True; r3.font.size=Pt(12); r3.font.color.rgb=GRIS

    p4=doc.add_paragraph(); p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r4=p4.add_run("academy.abedong.org"); r4.font.size=Pt(12); r4.font.color.rgb=VERT2; r4.underline=True

    doc.add_page_break()
    h("PARTIE 1 — 5 Posts LinkedIn prêts à publier",size=18)
    doc.add_paragraph("Posts basés sur les offres junior du jour. Copier-coller sur la page LinkedIn ABED Academy.").paragraph_format.space_after=Pt(12)

    for i,post in enumerate(posts_linkedin):
        sep()
        h(f"Post #{i+1} — {post.get('titre','')}",size=13,color=VERT2)
        p_c=doc.add_paragraph(post.get("contenu",""))
        p_c.paragraph_format.left_indent=Cm(0.8); p_c.paragraph_format.space_after=Pt(6)
        p_t=doc.add_paragraph()
        r_t=p_t.add_run(post.get("hashtags",""))
        r_t.font.color.rgb=RGBColor(0x18,0x5F,0xA5); r_t.font.size=Pt(9)
        p_t.paragraph_format.left_indent=Cm(0.8)
        doc.add_paragraph()

    doc.add_page_break()
    h("PARTIE 2 — 2 Scripts TikTok",size=18)
    doc.add_paragraph("Scripts pour vidéos 30-60 secondes basés sur les offres réelles du jour.").paragraph_format.space_after=Pt(12)

    for i,s in enumerate(scripts_tiktok):
        sep()
        h(f"TikTok #{i+1} — {s.get('titre','')}",size=13,color=ORANGE)
        p_a=doc.add_paragraph()
        ra=p_a.add_run("Accroche (3 premières sec) : "); ra.bold=True; ra.font.color.rgb=ORANGE; ra.font.size=Pt(10)
        p_a.add_run(s.get("accroche_visuelle","")).font.size=Pt(10)
        h("Texte parlé :",size=11,color=VERT)
        p_s=doc.add_paragraph(s.get("script_parle",""))
        p_s.paragraph_format.left_indent=Cm(0.8)
        p_te=doc.add_paragraph()
        rte=p_te.add_run("Texte écran : "); rte.bold=True; rte.font.size=Pt(10); rte.font.color.rgb=VERT
        p_te.add_run(s.get("texte_ecran","")).font.size=Pt(10)
        p_th=doc.add_paragraph()
        rth=p_th.add_run(s.get("hashtags",""))
        rth.font.color.rgb=RGBColor(0x18,0x5F,0xA5); rth.font.size=Pt(9)
        doc.add_paragraph()

    doc.add_page_break()
    h("ANNEXE — Récapitulatif des 10 offres",size=16)
    for i,o in enumerate(offers):
        p=doc.add_paragraph()
        rn=p.add_run(f"{i+1}. "); rn.bold=True; rn.font.color.rgb=VERT
        rt=p.add_run(o.get("titre","")); rt.bold=True; rt.font.size=Pt(10)
        p2=doc.add_paragraph(); p2.paragraph_format.left_indent=Cm(0.8)
        r_o=p2.add_run(f"{o.get('org','')} | {o.get('pays','')} | Exp: {o.get('mention_experience','?')} | ")
        r_o.font.size=Pt(9); r_o.font.color.rgb=GRIS
        r_p=p2.add_run(o.get("_parcours_matche",{}).get("nom",""))
        r_p.font.size=Pt(9); r_p.font.color.rgb=VERT2; r_p.bold=True
        p2.paragraph_format.space_after=Pt(4)

    doc.save(output_path)
    log.info(f"Word généré : {output_path}")
    return output_path


def send_daily_email(excel_path: Path, word_path: Path, offers: list[dict]):
    resend.api_key = os.environ.get("RESEND_API_KEY","")
    DESTINATAIRES  = ["olla.admi@gmail.com","adriendogo@gmail.com","prudencedogo@gmail.com","yetongnongbaguidi@gmail.com"]
    today     = date.today().strftime("%d/%m/%Y")
    nb_offres = len(offers)
    secteurs  = list(set(o.get("secteur","") for o in offers))[:4]

    corps_html = f"""
<html><body style="font-family:Arial,sans-serif;color:#2C2C2A;max-width:600px;margin:0 auto;">
<div style="background:#0F6E56;padding:24px;border-radius:8px 8px 0 0;text-align:center;">
  <h1 style="color:white;margin:0;font-size:22px;">ABED ACADEMY</h1>
  <p style="color:#E1F5EE;margin:4px 0 0;font-size:14px;">Bulletin Marketing Quotidien — {today}</p>
</div>
<div style="background:#fff;padding:24px;border:1px solid #E0DED8;">
  <p>Bonjour à toute l'équipe,</p>
  <p>Voici le bulletin marketing du <strong>{today}</strong> —
  <strong>{nb_offres} offres junior</strong> (max 3 ans d'exp) analysées
  dans les secteurs : <strong>{", ".join(secteurs)}</strong>.</p>
  <div style="background:#E1F5EE;padding:12px 16px;border-radius:6px;margin-bottom:12px;">
    <strong>📊 Excel</strong> — Top 10 offres + parcours ABED recommandé + liens directs
  </div>
  <div style="background:#FAEEDA;padding:12px 16px;border-radius:6px;margin-bottom:16px;">
    <strong>📝 Word</strong> — 5 posts LinkedIn + 2 scripts TikTok prêts à publier
  </div>
  <h3 style="color:#0F6E56;">🎯 Top 5 offres junior du jour</h3>
  <table style="width:100%;border-collapse:collapse;font-size:13px;">
    <tr style="background:#0F6E56;color:white;">
      <th style="padding:8px;text-align:left;">Poste</th>
      <th style="padding:8px;text-align:left;">Org / Pays</th>
      <th style="padding:8px;text-align:left;">Parcours ABED</th>
    </tr>
    {"".join(f'<tr style="background:{"#F8F7F2" if i%2==0 else "#fff"};"><td style="padding:8px;border-bottom:1px solid #E0DED8;"><strong>{o.get("titre","")[:55]}</strong></td><td style="padding:8px;border-bottom:1px solid #E0DED8;color:#5F5E5A;">{o.get("org","")[:30]} — {o.get("pays","")}</td><td style="padding:8px;border-bottom:1px solid #E0DED8;color:#0F6E56;font-size:11px;">{o.get("_parcours_matche",{}).get("nom","")[:40]}</td></tr>' for i,o in enumerate(offers[:5]))}
  </table>
  <div style="margin-top:20px;padding:14px;background:#E1F5EE;border-radius:6px;text-align:center;">
    <a href="https://academy.abedong.org" style="color:#0F6E56;font-weight:600;font-size:14px;">academy.abedong.org</a>
  </div>
</div>
<div style="background:#F8F7F2;padding:10px;text-align:center;font-size:11px;color:#888;">
  Généré automatiquement — {datetime.now().strftime('%d/%m/%Y à %H:%M')}
</div>
</body></html>"""

    attachments = []
    for path,fname,ctype in [
        (excel_path,f"ABED_Top10_Junior_{date.today().isoformat()}.xlsx",
         "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
        (word_path,f"ABED_Marketing_{date.today().isoformat()}.docx",
         "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    ]:
        if path.exists():
            with open(path,"rb") as f: content=base64.b64encode(f.read()).decode()
            attachments.append({"filename":fname,"content":content,"type":ctype})

    params = {
        "from":        "ABED Academy <noreply@abedong.org>",
        "to":          DESTINATAIRES,
        "subject":     f"[ABED Academy] Bulletin Marketing {today} — {nb_offres} offres junior",
        "html":        corps_html,
        "attachments": attachments,
    }
    try:
        response = resend.Emails.send(params)
        log.info(f"Email envoyé via Resend — id: {response['id']}")
    except Exception as e:
        log.error(f"Erreur Resend: {e}")
        raise


def run_marketing_pipeline(offers: list[dict], output_dir: Path):
    if not offers:
        log.warning("Aucune offre pour le pipeline marketing.")
        return
    log.info("=== PIPELINE MARKETING ABED ACADEMY ===")
    top10 = select_top10_for_abed(offers)
    log.info(f"Top 10 offres junior sélectionnées")
    top10 = enrich_offers_for_marketing(top10)
    output_dir.mkdir(parents=True, exist_ok=True)
    excel_path = output_dir / f"ABED_Top10_{date.today().isoformat()}.xlsx"
    word_path  = output_dir / f"ABED_Marketing_{date.today().isoformat()}.docx"
    generate_excel(top10, excel_path)
    generate_word(top10, word_path)
    try:
        send_daily_email(excel_path, word_path, top10)
    except Exception as e:
        log.error(f"Email non envoyé (fichiers disponibles dans exports/) : {e}")
    log.info("=== PIPELINE MARKETING TERMINÉ ===")
